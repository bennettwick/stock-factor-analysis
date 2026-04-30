#!/usr/bin/env python
"""
build_report.py — Generate Factor_Model_Report.docx.

Run: uv run python build_report.py
"""

import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── helpers ────────────────────────────────────────────────────────────────

def body(doc, text):
    """Justified 11pt paragraph."""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def placeholder(doc, label):
    """Centered gray italic placeholder line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'[ {label} ]')
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    return p


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


# ── document ───────────────────────────────────────────────────────────────

doc = Document()
doc.styles['Normal'].font.size = Pt(11)


# ══════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════

title_para = doc.add_heading(
    'Factor Model and Time Series Analysis\nof Tesla (TSLA) and Disney (DIS)',
    level=0
)
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

for label in ['[YOUR NAME]', '[COURSE NAME / NUMBER]', '[INSTRUCTOR NAME]']:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(datetime.date.today().strftime('%B %d, %Y'))
r.font.size = Pt(12)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════

h1(doc, '1. Introduction')

body(doc, (
    'This report analyzes the equity risk profiles of Tesla, Inc. (TSLA, Consumer Discretionary '
    '/ Electric Vehicles) and The Walt Disney Company (DIS, Communication Services / Entertainment) '
    'over the fifteen-year period from January 3, 2011 through December 31, 2025 — 3,771 trading '
    'days of daily return data. These two stocks were selected to juxtapose a high-volatility, '
    'high-growth speculative name against a mature, asset-intensive media and entertainment '
    'conglomerate, a pairing that spans much of the systematic-risk spectrum observable among '
    'large-cap U.S. equities. TSLA began the sample period as a small-cap electric vehicle '
    'startup with minimal revenues and grew into one of the largest companies in the world by '
    'market capitalization; DIS entered the sample as an established global entertainment leader '
    'and navigated structural disruptions in linear television, the launch of the Disney+ streaming '
    'platform, and the near-complete shutdown of its theme-park business during the COVID-19 '
    'pandemic. The contrast between the two stocks — a high-beta, speculative-growth name versus '
    'a market-neutral, value-tilted conglomerate — motivates the comparative framework used '
    'throughout this analysis.'
))

body(doc, (
    'The analysis proceeds in four stages. First, both return series are characterized using '
    'time-series methods — price and return plots, rolling 60-day statistics, Augmented '
    'Dickey-Fuller stationarity tests, and ACF/PACF analysis of raw and squared returns — to '
    'establish their statistical properties before any factor model is fit. These diagnostics '
    'confirm that prices are I(1) non-stationary while returns are I(0) stationary, validating '
    'OLS regression in return space, and reveal pronounced ARCH-type volatility clustering in '
    'both series. Second, the Capital Asset Pricing Model (CAPM) is estimated for each stock, '
    'yielding unconditional estimates of market sensitivity (beta) and risk-adjusted abnormal '
    'return (alpha). Third, the Fama-French three-factor model (FF3) is estimated, adding the '
    'SMB (Small Minus Big) and HML (High Minus Low) factors to disentangle size and value '
    'exposures from pure market risk; multicollinearity is assessed via Variance Inflation '
    'Factors and model improvement is assessed via adjusted R-squared. Fourth, rolling '
    '252-day CAPM betas are computed to assess whether unconditional estimates are stable '
    'across the sample period, and FF3 residuals are subjected to stationarity and '
    'autocorrelation diagnostics to characterize what the linear model leaves unexplained.'
))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# 2. METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════

h1(doc, '2. Methodology')

body(doc, (
    'Daily adjusted close prices for TSLA and DIS were downloaded from Yahoo Finance for '
    'January 3, 2011 through December 31, 2025 (end date set to January 1, 2026 to include '
    'December 31, 2025, since the yfinance API end date is exclusive). Simple daily returns '
    'were computed as r_t = (P_t − P_{t−1}) / P_{t−1} using the pct_change() method, '
    'producing decimal-scale returns (e.g., 0.005 for a 0.5% gain). Fama-French three-factor '
    'data (Mkt-RF, SMB, HML) and the daily risk-free rate (RF) were sourced from the provided '
    'fda1_stock_factor_data.csv file. A critical unit-alignment step is required: the FF data '
    'is stored in percent units (e.g., 0.50 for 0.50% return), while yfinance returns are in '
    'decimal units (0.005 for 0.50%). All four FF columns were therefore divided by 100 '
    'immediately after loading. Failing to apply this conversion would inflate FF factor '
    'values by a factor of 100 relative to the dependent variable, silently biasing every '
    'alpha and beta estimate by a corresponding factor. The two datasets were inner-joined '
    'on date index and residual NaN observations were dropped, yielding a clean '
    '3,771-observation panel spanning January 4, 2011 through December 31, 2025.'
))

body(doc, (
    'Excess returns were computed as r_i_excess = r_i − RF for each stock. The market '
    'excess return Mkt-RF was used directly from the FF file without further subtraction of '
    'RF, since it is already expressed as an excess return. CAPM was estimated as '
    'r_i − r_f = α + β(Mkt-RF) + ε via OLS with an intercept added using '
    'sm.add_constant(); the FF3 model extended this to include SMB and HML as additional '
    'regressors. All regressions were implemented with statsmodels OLS. VIFs were computed '
    'using the variance_inflation_factor function from statsmodels with the constant included '
    'in the predictor matrix, matching the linear space in which multicollinearity affects '
    'estimates. Rolling 252-day betas were computed using the covariance-variance '
    'formulation β_t = Cov(r_i_ex, Mkt-RF; 252) / Var(Mkt-RF; 252) via pandas rolling '
    'methods. Stationarity was assessed using the Augmented Dickey-Fuller test '
    '(statsmodels adfuller) and autocorrelation structure was examined via ACF and PACF '
    'plots at 40 lags.'
))

body(doc, (
    'A key assumption throughout is OLS linearity: the conditional expected excess return '
    'is modeled as a linear function of contemporaneous factor realizations. Given the '
    'well-documented ARCH effects confirmed by the squared-return ACF analysis, default OLS '
    'standard errors are likely downward-biased in high-volatility regimes. Heteroskedasticity-'
    'robust standard errors (HC3 or Newey-West) would be preferred for formal hypothesis '
    'testing on factor loadings; the regressions as presented use nonrobust standard errors '
    'per the course specification, and this limitation is acknowledged where it affects '
    'interpretation of significance levels.'
))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# 3. RESULTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, '3. Results')

# ── 3.1 Time Series Characterization ──────────────────────────────────────

h2(doc, '3.1 Time Series Characterization')

body(doc, (
    'Price levels for both stocks display the hallmark of non-stationarity: both drift '
    'persistently over the sample with no fixed mean, and TSLA\'s price grew by several '
    'orders of magnitude from its 2011 levels near $5 to peaks above $400. ADF tests confirm '
    'the visual impression: neither price series rejects the unit-root null (TSLA: ADF = '
    '−0.496, p = 0.893; DIS: ADF = −2.053, p = 0.264), while both return series reject '
    'it with ADF statistics far below critical values and p-values of essentially zero. '
    'Returns are I(0) stationary, confirming that working in return space is appropriate '
    'and that OLS regression of returns on factor returns will not produce spurious results '
    'driven by shared stochastic trends.'
))

placeholder(doc, 'INSERT FIGURE: Adjusted Close Prices — TSLA and DIS, Jan 2011–Dec 2025 (notebook Section 2.1)')
placeholder(doc, 'INSERT FIGURE: Daily Simple Returns — TSLA and DIS, Jan 2011–Dec 2025 (notebook Section 2.1)')
placeholder(doc, 'INSERT TABLE: ADF Test Results — Price Levels and Returns for TSLA and DIS (notebook Section 2.3)')

body(doc, (
    'Rolling 60-day statistics reveal two qualitatively distinct volatility regimes. The '
    'COVID-19 shock in March 2020 produced acute spikes in both stocks: TSLA\'s 60-day '
    'rolling annualized volatility briefly exceeded 150%, while DIS spiked above 80% — '
    'both far above their respective pre-pandemic baselines. The spike was sudden and '
    'short-lived, reflecting indiscriminate liquidation followed by a rapid V-shaped recovery. '
    'DIS\'s spike was amplified by the near-simultaneous closure of its theme parks, '
    'a direct revenue disruption layered on top of the market-wide panic. The 2022 Federal '
    'Reserve rate-hiking cycle produced a qualitatively different regime: sustained elevated '
    'volatility for TSLA throughout 2022, at roughly twice its post-COVID baseline, consistent '
    'with the compression of its long-duration growth-stock valuation under rising discount '
    'rates. DIS experienced a lower level of volatility elevation in 2022 and for a shorter '
    'duration, reflecting its more near-term cash-flow profile and its lesser sensitivity '
    'to discount-rate changes relative to a long-duration growth stock.'
))

placeholder(doc, 'INSERT FIGURE: 60-Day Rolling Mean and Annualized Volatility — TSLA and DIS (notebook Section 2.2)')

body(doc, (
    'ACF and PACF plots of raw daily returns show at most isolated, marginally significant '
    'spikes, with all remaining autocorrelations well within 95% confidence bands — '
    'consistent with weak-form market efficiency at the daily frequency. Neither stock '
    'exhibits exploitable return autocorrelation. By contrast, ACF plots of squared returns '
    'display statistically significant, slowly decaying autocorrelation across all 40 lags '
    'for both stocks: the canonical fingerprint of volatility clustering (ARCH effects). '
    'The magnitude of returns is serially correlated even when the level and sign are not. '
    'TSLA\'s squared-return autocorrelations decay more slowly than DIS\'s, indicating more '
    'persistent volatility shocks consistent with its higher sensitivity to speculative '
    'sentiment and headline risk.'
))

placeholder(doc, 'INSERT FIGURE: ACF and PACF of Daily Returns, 40 Lags — TSLA and DIS (notebook Section 2.4)')
placeholder(doc, 'INSERT FIGURE: ACF of Squared Daily Returns, 40 Lags — TSLA and DIS (notebook Section 2.4)')

# ── 3.2 CAPM ──────────────────────────────────────────────────────────────

h2(doc, '3.2 CAPM / Market Model')

body(doc, (
    'The Capital Asset Pricing Model was estimated over the full 3,771-observation sample '
    'for each stock. Key results are presented in Table 1.'
))

placeholder(doc, 'INSERT TABLE: CAPM Summary — Alpha, Beta, p-values, R-squared for TSLA and DIS (notebook Section 3.3)')

body(doc, (
    'TSLA exhibits a market beta of 1.554, well above unity, classifying it as an aggressive '
    'stock that amplifies broad market movements: a 1% market gain is associated on average '
    'with a 1.554% TSLA gain. TSLA\'s alpha is statistically significant at the 5% level '
    '(α = 0.0012 per day, p = 0.015) — its realized daily returns exceeded CAPM '
    'predictions over the full sample, driven primarily by its extraordinary 2019–2021 '
    'appreciation. This alpha is a backward-looking, full-sample estimate and should not '
    'be interpreted as a signal of repeatable future outperformance. The market factor '
    'explains 23.3% of TSLA\'s daily return variance; the remaining 76.7% is idiosyncratic, '
    'reflecting TSLA-specific dynamics that a single broad factor cannot capture.'
))

body(doc, (
    'DIS exhibits a market beta of 0.989, essentially exactly unity, confirming that it '
    'tracks broad market movements at a one-for-one rate. DIS\'s alpha is not statistically '
    'significant (α = −0.0001 per day, p = 0.622), consistent with efficient market '
    'pricing for a mature large-cap conglomerate. The CAPM explains 43.9% of DIS\'s daily '
    'return variance, a substantially higher fraction than TSLA\'s, indicating that DIS\'s '
    'returns are more tightly coupled to the broad market factor with proportionally less '
    'idiosyncratic noise.'
))

body(doc, (
    'Residual diagnostics (time-series plot, histogram, Q-Q plot) reveal that both stocks\' '
    'CAPM residuals violate OLS homoskedasticity: time-series plots show pronounced volatility '
    'clustering, histograms are highly leptokurtic relative to the normal distribution, and '
    'Q-Q plots display systematic fat-tail deviations at both ends. These pathologies do not '
    'bias the beta estimates but render default standard errors unreliable, particularly '
    'in high-volatility sub-periods.'
))

placeholder(doc, 'INSERT FIGURE: CAPM Residual Diagnostics (Residuals over Time, Histogram, Q-Q Plot) — TSLA and DIS (notebook Section 3.5)')

# ── 3.3 FF3 ───────────────────────────────────────────────────────────────

h2(doc, '3.3 Fama-French Three-Factor Model')

body(doc, (
    'The Fama-French three-factor model was estimated by adding SMB and HML to the CAPM '
    'regression for each stock. Factor loadings, p-values, and goodness-of-fit statistics '
    'are presented in Tables 2–4.'
))

placeholder(doc, 'INSERT TABLE: FF3 Factor Loadings, p-values, R-squared, Adj. R-squared — TSLA and DIS (notebook Section 4.2)')
placeholder(doc, 'INSERT TABLE: VIF — Fama-French Factor Multicollinearity Check (notebook Section 4.4)')
placeholder(doc, 'INSERT TABLE: CAPM vs. FF3 R-squared Comparison — TSLA and DIS (notebook Section 4.5)')

body(doc, (
    'VIF values for all three factors are below 1.14 (Mkt-RF: 1.076, SMB: 1.139, HML: 1.074), '
    'confirming that multicollinearity is negligible and each loading reflects that factor\'s '
    'unique contribution to explaining excess returns.'
))

body(doc, (
    'For TSLA, the FF3 market beta (1.448) is slightly lower than the CAPM estimate (1.555), '
    'with a portion of the loading reallocated to the SMB factor once size is controlled for. '
    'The SMB loading is positive and highly significant (+0.541, p < 0.001): rather than the '
    'negative sign typical of established large-caps, TSLA\'s returns co-move with small-cap '
    'stocks, reflecting its genuine small/mid-cap status during the early sample years '
    '(2011–2017) and its sustained high-beta-to-risk-sentiment profile that causes it to '
    'behave more like speculative growth stocks than large-cap value incumbents regardless '
    'of market capitalization. The HML loading is strongly negative and significant '
    '(−0.909, p < 0.001), the definitive signature of a growth stock priced on future '
    'earnings expectations rather than current book value. Even after three-factor adjustment, '
    'TSLA\'s alpha remains significant at the 1% level (0.0013/day, p = 0.009), reflecting '
    'its extraordinary realized appreciation over the full sample. FF3 R-squared rises to '
    '27.2%, a 3.9 percentage-point improvement over CAPM that is confirmed by rising '
    'adjusted R-squared.'
))

body(doc, (
    'For DIS, the FF3 market beta is essentially unchanged at 0.990. The SMB loading is '
    'near zero and not statistically significant (+0.046, p = 0.171), consistent with DIS\'s '
    'stable large-cap classification. The HML loading is positive and highly significant '
    '(+0.229, p < 0.001): Disney\'s substantial tangible assets — theme parks, resort '
    'real estate, physical production infrastructure, and long-lived intellectual property '
    '— give it a higher book-to-market ratio than capital-light technology companies, '
    'producing the positive HML loading characteristic of value-tilted large-caps. This '
    'value characteristic distinguishes DIS sharply from TSLA\'s anti-value profile. FF3 '
    'R-squared rises to 45.2% (+1.3 percentage points over CAPM), driven primarily by '
    'the significant HML factor. DIS\'s alpha remains insignificant after three-factor '
    'adjustment (p = 0.667).'
))

placeholder(doc, 'INSERT FIGURE: FF3 Residual Diagnostics (Residuals over Time, Histogram, Q-Q Plot) — TSLA and DIS (notebook Section 4.6)')

# ── 3.4 Synthesis ─────────────────────────────────────────────────────────

h2(doc, '3.4 Synthesis: Rolling Beta and Residual Diagnostics')

body(doc, (
    'Rolling 252-day CAPM betas were computed via the covariance-variance formulation '
    '(β_t = Cov(r_i_ex, Mkt-RF; 252) / Var(Mkt-RF; 252)), capturing how each stock\'s '
    'systematic market sensitivity evolved over time. The first 251 observations produce '
    'NaN values by construction and are excluded from the chart.'
))

placeholder(doc, 'INSERT FIGURE: Rolling 252-Day CAPM Beta — TSLA and DIS with β=1 Reference Line (notebook Section 5.1)')

body(doc, (
    'TSLA\'s rolling beta is highly unstable, spanning a wide range from sub-1.0 readings '
    'in the early sample to values well above 2.0 during the 2020–2021 speculative rally. '
    'The COVID-19 shock in March 2020 caused a transient spike in both stocks\' rolling betas '
    '— the standard crisis-period convergence where cross-asset correlations compress and '
    'all betas temporarily migrate toward the market. The 2022 Federal Reserve rate-hiking '
    'cycle kept TSLA\'s beta elevated and volatile, as the compression of long-horizon earnings '
    'expectations amplified TSLA\'s drawdown relative to the broad market. DIS\'s rolling '
    'beta is considerably more stable, tracking near 1.0 for most of the sample with only '
    'transient deviations during acute stress; the two betas diverge persistently from 2020 '
    'onward, confirming that the higher full-sample TSLA beta reflects a durably different '
    'risk profile rather than a handful of outlier episodes.'
))

body(doc, (
    'ADF tests on FF3 residuals reject the unit-root null at p ≈ 0.000 for both stocks '
    '(TSLA: ADF = −34.94; DIS: ADF = −35.82), confirming that the FF3 model '
    'produces stationary residuals as expected when stationary returns are regressed on '
    'stationary factors. No structural misspecification drives an integrated residual process.'
))

placeholder(doc, 'INSERT TABLE: ADF Test on FF3 Residuals — TSLA and DIS (notebook Section 5.3)')

body(doc, (
    'ACF and PACF plots of FF3 residuals in levels show at most isolated marginally significant '
    'spikes, broadly consistent with white noise: the three factors adequately capture the '
    'conditional mean structure of daily excess returns. By contrast, ACF plots of squared '
    'FF3 residuals display highly significant, slowly decaying autocorrelation across all '
    '40 lags for both stocks — a pattern nearly identical to the squared raw return ACFs. '
    'Removing the Fama-French factors does not materially reduce volatility clustering: the '
    'factors explain conditional mean dynamics but leave conditional variance dynamics '
    'entirely unmodeled. ARCH effects are primarily idiosyncratic, pointing to GARCH-family '
    'models as the natural extension for a complete characterization of daily return risk.'
))

placeholder(doc, 'INSERT FIGURE: ACF and PACF of FF3 Residuals, 40 Lags — TSLA and DIS (notebook Section 5.4)')
placeholder(doc, 'INSERT FIGURE: ACF of Squared FF3 Residuals, 40 Lags — TSLA and DIS (notebook Section 5.4)')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════

h1(doc, '4. Discussion')

# ── 4a ────────────────────────────────────────────────────────────────────

h2(doc, '4a. Market Sensitivity: CAPM Beta and Rolling Beta')

body(doc, (
    'TSLA is unambiguously the more market-sensitive stock across every metric examined. '
    'Its full-sample CAPM beta of 1.554 is 56 percentage points above DIS\'s beta of 0.989, '
    'meaning TSLA amplifies a 1% market move by approximately 1.554% while DIS merely '
    'tracks it at 0.989%. The FF3 model slightly reduces TSLA\'s market beta to 1.448 as '
    'a portion of the sensitivity is attributed to the SMB factor, but the gap with DIS '
    'remains large (1.448 vs. 0.990). The rolling beta analysis confirms this difference '
    'and adds a temporal dimension: TSLA\'s rolling beta is highly unstable, spanning a '
    'wide range from below 1.0 in the early sample to well above 2.0 at peak moments, '
    'while DIS\'s rolling beta tracks steadily near 1.0. The practical implication is that '
    'a portfolio manager using TSLA\'s full-sample beta as a hedge ratio would face '
    'substantial tracking error in most sub-periods, as the realized beta at any given '
    'time may differ significantly from the long-run unconditional estimate. DIS\'s beta '
    'stability makes its unconditional estimate a far more reliable summary of its actual '
    'systematic risk exposure in any given period.'
))

# ── 4b ────────────────────────────────────────────────────────────────────

h2(doc, '4b. SMB and HML Loadings vs. Prior Expectations')

body(doc, (
    'TSLA\'s HML loading meets prior expectations precisely: its strongly negative value '
    '(−0.909, p < 0.001) confirms TSLA as a canonical growth stock, priced far above '
    'book value on expectations of future earnings growth and co-moving with other low '
    'book-to-market companies. Its SMB loading (+0.541, p < 0.001), however, initially '
    'contradicts naive intuition but makes economic sense on closer examination. For much '
    'of the early sample (2011–2017), TSLA was genuinely a small/mid-cap company. Even '
    'after becoming a mega-cap, its speculative return profile — high sensitivity to '
    'risk appetite, narrative-driven trading, and retail investor participation — causes '
    'it to co-move with high-risk small-cap growth stocks rather than with the large-cap '
    'value incumbents that carry negative SMB loadings. Factor loadings measure return '
    'co-movement, not simply size classification.'
))

body(doc, (
    'DIS\'s factor loadings also match expectations. Its positive and significant HML '
    'loading (+0.229, p < 0.001) correctly identifies DIS as a value-tilted stock: '
    'Disney\'s substantial tangible assets — theme parks, resort real estate, physical '
    'production infrastructure, and long-lived intellectual property — produce a higher '
    'book-to-market ratio than capital-light technology companies, and the positive HML '
    'loading captures this value characteristic. DIS\'s SMB loading is near zero and '
    'insignificant (+0.046, p = 0.171), consistent with stable large-cap classification '
    'and no systematic small-cap tilt. Both stocks\' factor profiles are economically '
    'coherent once the SMB anomaly for TSLA is interpreted as a return-comovement '
    'signature rather than a size misclassification.'
))

# ── 4c ────────────────────────────────────────────────────────────────────

h2(doc, '4c. Did FF3 Materially Improve Over CAPM?')

body(doc, (
    'For both stocks, CAPM does the dominant explanatory work. The market factor alone '
    'accounts for 23.3% of TSLA\'s and 43.9% of DIS\'s daily return variance. Adding SMB '
    'and HML raises R-squared by 3.9 percentage points for TSLA (to 27.2%) and 1.3 '
    'percentage points for DIS (to 45.2%); both gains are confirmed by rising adjusted '
    'R-squared, confirming genuine incremental value rather than mechanical overfitting. '
    'In relative terms, FF3 provides a larger proportional improvement for TSLA '
    '(approximately +17% relative to CAPM R-squared) than for DIS (approximately +3%). '
    'For TSLA, both SMB and HML are significant contributors to the improvement; for DIS, '
    'the gain is driven almost entirely by the significant HML factor, while the '
    'insignificant SMB adds negligible fit. The conclusion is that CAPM captures the '
    'dominant systematic risk for both stocks, while FF3 provides meaningful but modest '
    'incremental precision — particularly for TSLA, where the size and value dimensions '
    'represent economically distinct exposures not subsumed by the market factor.'
))

# ── 4d ────────────────────────────────────────────────────────────────────

h2(doc, '4d. Are FF3 Residuals Close to White Noise?')

body(doc, (
    'In levels, the FF3 residuals are broadly consistent with white noise: ACF and PACF '
    'plots show at most isolated marginally significant spikes, and ADF tests confirm '
    'stationarity with p ≈ 0.000 for both stocks. The linear factor model adequately '
    'captures the conditional mean structure of daily excess returns, and there is no '
    'evidence of systematic residual autocorrelation in levels. In variance, however, '
    'significant structure remains: ACF plots of squared FF3 residuals display highly '
    'significant, slowly decaying autocorrelation across all 40 lags for both TSLA and '
    'DIS — a pattern nearly identical in character to the squared raw return ACFs from '
    'the time-series characterization section. Removing the three Fama-French factors '
    'does not materially reduce volatility clustering: the factors explain conditional '
    'mean return dynamics but leave conditional variance dynamics entirely unmodeled. '
    'ARCH effects are primarily idiosyncratic rather than factor-driven. The practical '
    'implication is that FF3 is a good model for the conditional mean of daily returns '
    'but a poor model for conditional risk, requiring GARCH-family extensions for accurate '
    'characterization of time-varying uncertainty.'
))

# ── 4e ────────────────────────────────────────────────────────────────────

h2(doc, '4e. Did Rolling Beta Change Meaningfully Over Time?')

body(doc, (
    'For TSLA, rolling beta changed dramatically over the sample period, with distinct '
    'regime-linked patterns. In the early years (2011–2013), TSLA was a small-cap '
    'startup with modest market correlation and relatively moderate rolling beta. Beta '
    'rose through the 2013–2014 retail enthusiasm period following the Model S commercial '
    'launch. The COVID-19 shock in March 2020 caused a sharp transient spike in both '
    'stocks\' rolling betas — the standard crisis-period effect where diversification '
    'disappears, cross-asset correlations converge, and all risky assets become temporarily '
    'more correlated with the market. The 2020–2021 speculative growth rally, fueled by '
    'near-zero interest rates and retail participation, pushed TSLA\'s rolling beta to some '
    'of its highest realized values as high-multiple assets moved in concert. The 2022 '
    'Federal Reserve rate-hiking cycle then kept TSLA\'s beta elevated and volatile as '
    'rate sensitivity drove outsized negative returns relative to the broad market. '
    'For DIS, rolling beta evolved far less dramatically, remaining near 1.0 for most of '
    'the sample; the March 2020 COVID spike represents the single most notable deviation, '
    'amplified by the real-time closure of its theme parks on top of the market-wide '
    'dislocation. The two stocks\' rolling betas diverge persistently from 2020 onward, '
    'reflecting their fundamentally different sensitivities to the monetary policy and '
    'sentiment regime that defined the post-pandemic period.'
))

# ── 4f ────────────────────────────────────────────────────────────────────

h2(doc, '4f. Limitations of Full-Sample Linear Models')

body(doc, (
    'Full-sample linear factor models face two core limitations when applied to daily equity '
    'data over long horizons with shifting market regimes. First, the unconditional beta '
    'assumes parameter stability: a single coefficient estimated over 15 years implicitly '
    'treats TSLA\'s beta in 2011 (when it was a small-cap startup) as equivalent to its beta '
    'in 2024 (when it was a mega-cap growth stock in the top 10 of the S&P 500). The rolling '
    'beta analysis demonstrates that this assumption is materially false for TSLA — the '
    'full-sample beta of 1.554 is an average over a highly variable time-series of '
    'instantaneous betas, masking substantial regime-dependent variation. For DIS, the '
    'stability assumption is closer to valid, but even there the COVID episode demonstrates '
    'that transient beta instability can arise during systemic market stress. Investors and '
    'risk managers who use full-sample CAPM/FF3 betas as forward-looking estimates implicitly '
    'assume that the past distribution of market regimes will persist, an assumption that '
    'is frequently violated at precisely the moments when accurate risk estimates are most '
    'consequential.'
))

body(doc, (
    'Second, linear factor models explain conditional mean returns but not conditional '
    'variance. The persistent ARCH effects in FF3 residuals for both stocks confirm that '
    'the variance of idiosyncratic returns follows a dynamic process — volatility clusters, '
    'and the clustering is unrelated to the Fama-French factors. OLS standard errors '
    'estimated under the homoskedasticity assumption are invalid in high-volatility '
    'regimes: they understate parameter uncertainty in periods like March 2020 and overstate '
    'it in calm periods. As descriptive tools applied to historical data, CAPM and FF3 '
    'perform well — they identify economically meaningful and statistically robust '
    'differences in risk exposure between TSLA and DIS. As forecasting tools, they are '
    'more limited: the betas they estimate are conditional on the historical sample '
    'distribution of market regimes, and structural changes in the composition or '
    'investor base of a stock can render those estimates stale. A complete risk model '
    'for either stock would require either time-varying-parameter betas (e.g., Kalman '
    'filter or rolling regression) or GARCH-family extensions to characterize the '
    'time-varying conditional variance that static linear models leave behind.'
))

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════

h1(doc, '5. Conclusion')

body(doc, (
    'This analysis characterized the systematic risk profiles of Tesla (TSLA) and Disney '
    '(DIS) over the January 2011 – December 2025 period using CAPM, the Fama-French '
    'three-factor model, and supporting time-series diagnostics. The two stocks occupy '
    'sharply different positions on the risk spectrum: TSLA is a high-beta, speculative-'
    'growth stock with a strongly negative HML loading, positive SMB loading reflecting '
    'its small-cap origins and speculative return co-movement, highly unstable rolling '
    'beta, and statistically significant positive alpha that reflects its extraordinary '
    'realized appreciation over the full sample period; DIS is a market-neutral, value-'
    'tilted large-cap conglomerate with a stable near-unity beta, positive HML exposure '
    'rooted in its tangible asset holdings, and no detectable alpha after three-factor '
    'adjustment. Both stocks exhibit persistent ARCH-type volatility clustering in their '
    'FF3 residuals, confirming that the three-factor linear model fully characterizes the '
    'mean structure of daily excess returns but leaves conditional variance dynamics '
    'unmodeled — a structural limitation that motivates GARCH-family extensions for '
    'any application requiring accurate time-varying risk characterization. Together, '
    'TSLA and DIS illustrate how the Fama-French framework captures economically meaningful '
    'cross-sectional differences in risk exposure while also exposing the fundamental '
    'boundary of what static linear models can describe about the dynamic behavior of '
    'daily equity returns.'
))


# ══════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════

doc.save('Factor_Model_Report.docx')
print('Saved Factor_Model_Report.docx')
print()

FIGURES = [
    ('2.1', 'Adjusted Close Prices — TSLA and DIS'),
    ('2.1', 'Daily Simple Returns — TSLA and DIS'),
    ('2.2', '60-Day Rolling Mean and Annualized Volatility'),
    ('2.4', 'ACF and PACF of Daily Returns (40 lags)'),
    ('2.4', 'ACF of Squared Daily Returns (40 lags)'),
    ('3.5', 'CAPM Residual Diagnostics (2x3 grid: time series, histogram, Q-Q)'),
    ('4.6', 'FF3 Residual Diagnostics (2x3 grid: time series, histogram, Q-Q)'),
    ('5.1', 'Rolling 252-Day CAPM Beta with beta=1 reference line'),
    ('5.4', 'ACF and PACF of FF3 Residuals (40 lags)'),
    ('5.4', 'ACF of Squared FF3 Residuals (40 lags)'),
]

TABLES = [
    ('2.3', 'ADF Test Results — price levels and returns'),
    ('3.3', 'CAPM Summary — alpha, beta, p-values, R^2'),
    ('4.2', 'FF3 Summary — all factor loadings, p-values, R^2, adj. R^2'),
    ('4.4', 'VIF — Fama-French factor multicollinearity'),
    ('4.5', 'CAPM vs. FF3 R^2 Comparison'),
    ('5.3', 'ADF Test on FF3 Residuals'),
]

print('=== Placeholders requiring manual completion ===')
print()
print(f'FIGURES ({len(FIGURES)} total — export from notebook and paste into Word):')
for sec, desc in FIGURES:
    print(f'  notebook {sec}: {desc}')
print()
print(f'TABLES ({len(TABLES)} total — paste formatted version from notebook):')
for sec, desc in TABLES:
    print(f'  notebook {sec}: {desc}')
print()
print('TITLE PAGE (3 items):')
print('  [YOUR NAME]')
print('  [COURSE NAME / NUMBER]')
print('  [INSTRUCTOR NAME]')
print()
print('All prose is complete — no other placeholders.')
